"""
DOCX Parser for Personal AI System

Extracts text content from DOCX files for RAG embedding
"""

from pathlib import Path
from typing import List
import re

try:
    from docx import Document
except ImportError:
    Document = None


class DOCXParser:
    """Parser for DOCX files."""
    
    def __init__(self):
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )
    
    def parse(self, file_path: Path) -> str:
        """
        Parse DOCX file and extract text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Text content extracted from DOCX
        """
        if Document is None:
            raise ImportError("python-docx not installed")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {str(e)}")
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a table in the document."""
        rows_text = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                rows_text.append(" | ".join(cells_text))
        return "\n".join(rows_text) if rows_text else ""
    
    def extract_structure(self, file_path: Path) -> dict:
        """Extract document structure information."""
        if Document is None:
            raise ImportError("python-docx not installed")
        
        doc = Document(file_path)
        
        return {
            'paragraphs': len(list(doc.paragraphs)),
            'tables': len(list(doc.tables)),
            'sections': len(doc.sections),
        }
